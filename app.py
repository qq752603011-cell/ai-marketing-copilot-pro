import streamlit as st
from openai import OpenAI
from datetime import datetime
import requests
import tempfile
import os
import json
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment
from docx import Document

# =========================
# 页面配置
# =========================

st.set_page_config(
    page_title="AI Marketing Copilot V3",
    page_icon="🚀",
    layout="wide"
)

if "history" not in st.session_state:
    st.session_state.history = []

# =========================
# 小红书封面图生成
# 用户输入提示词，硅基流动 Kolors 模型生成
# =========================

# 小红书封面常用场景提示词模板
COVER_TEMPLATES = {
    "产品平铺": "flat lay product photography, {product} on white marble background, minimalist aesthetic, soft natural lighting, xiaohongshu style, high quality",
    "生活场景": "lifestyle photography, {product} in cozy home environment, warm natural light, person using product, xiaohongshu aesthetic, commercial photo",
    "简约纯色": "product shot, {product}, solid pastel background, centered composition, clean minimalist style, studio lighting, xiaohongshu cover",
    "户外场景": "outdoor lifestyle, {product}, natural scenery background, golden hour light, authentic feel, xiaohongshu style",
    "礼盒展示": "{product} gift packaging, luxury unboxing, white and gold tones, elegant presentation, xiaohongshu cover photo",
}

COVER_MODELS = [
    "Kwai-Kolors/Kolors",
    "stabilityai/stable-diffusion-xl-base-1.0",
    "stabilityai/stable-diffusion-3-medium",
]

def generate_cover_image(user_prompt):
    """
    用户自定义提示词生成小红书封面
    依次尝试多个模型，第一个成功即返回
    """
    sf_key = st.secrets.get("SILICONFLOW_API_KEY", "")
    if not sf_key:
        return None, "未配置 SILICONFLOW_API_KEY"

    for model in COVER_MODELS:
        try:
            resp = requests.post(
                "https://api.siliconflow.cn/v1/images/generations",
                headers={
                    "Authorization": f"Bearer {sf_key}",
                    "Content-Type": "application/json"
                },
                json={
                    "model": model,
                    "prompt": user_prompt,
                    "image_size": "768x1024",
                    "num_inference_steps": 20,
                    "batch_size": 1
                },
                timeout=90
            )
            result = resp.json()

            if result.get("images"):
                img_url = result["images"][0]["url"]
                img_resp = requests.get(img_url, timeout=30)
                if img_resp.status_code == 200:
                    cover_path = f"cover_{datetime.now().strftime('%H%M%S')}.jpg"
                    with open(cover_path, "wb") as f:
                        f.write(img_resp.content)
                    return cover_path, None
            else:
                msg = result.get("message", str(result))
                if "disabled" in msg.lower() or "not exist" in msg.lower():
                    continue  # 换下一个模型
                return None, f"生成失败（{model}）：{msg}"

        except Exception as e:
            continue

    return None, "所有模型均不可用，请检查硅基流动账户余额"


# =========================
# 飞书工具函数
# =========================

def get_feishu_token():
    resp = requests.post(
        "https://open.feishu.cn/open-apis/auth/v3/tenant_access_token/internal",
        json={
            "app_id": st.secrets["FEISHU_APP_ID"],
            "app_secret": st.secrets["FEISHU_APP_SECRET"]
        }
    ).json()
    return resp.get("tenant_access_token")


def upload_image_to_feishu(image_path, token, app_token):
    url = "https://open.feishu.cn/open-apis/drive/v1/medias/upload_all"
    headers = {"Authorization": f"Bearer {token}"}
    file_name = os.path.basename(image_path)
    file_size = os.path.getsize(image_path)

    with open(image_path, "rb") as f:
        resp = requests.post(
            url,
            headers=headers,
            files={"file": (file_name, f, "image/jpeg")},
            data={
                "file_name": file_name,
                "parent_type": "bitable_file",
                "parent_node": app_token,
                "size": str(file_size)
            }
        )

    result = resp.json()
    if result.get("code") == 0:
        return result["data"]["file_token"]
    else:
        st.warning(f"图片上传失败 code={result.get('code')}：{result.get('msg')}")
        return None


def save_to_feishu(parsed, image_path=None, cover_path=None):
    token = get_feishu_token()
    base_id = st.secrets["BASE_ID"]
    table_id = st.secrets["TABLE_ID"]

    def to_str(v):
        if isinstance(v, list):
            return "\n".join(str(i) for i in v)
        return str(v) if v else ""

    fields = {
        "商品标题": to_str(parsed.get("商品标题", "")),
        "封面标题": to_str(parsed.get("封面标题", "")),
        "核心卖点": to_str(parsed.get("核心卖点", "")),
        "小红书文案": to_str(parsed.get("小红书文案", "")),
        "短视频口播稿": to_str(parsed.get("短视频口播稿", "")),
        "亚马逊卖点": to_str(parsed.get("亚马逊卖点", "")),
    }

    if image_path and os.path.exists(image_path):
        ft = upload_image_to_feishu(image_path, token, base_id)
        if ft:
            fields["产品图片"] = [{"file_token": ft}]

    if cover_path and os.path.exists(cover_path):
        ft = upload_image_to_feishu(cover_path, token, base_id)
        if ft:
            fields["小红书封面"] = [{"file_token": ft}]

    resp = requests.post(
        f"https://open.feishu.cn/open-apis/bitable/v1/apps/{base_id}/tables/{table_id}/records",
        headers={
            "Authorization": f"Bearer {token}",
            "Content-Type": "application/json"
        },
        json={"fields": fields}
    )
    return resp.json()


# =========================
# AI 结构化内容生成
# =========================

def generate_structured_content(product_name, target_customer, platform, features, client, model_name):
    prompt = f"""你是资深电商运营专家。请为以下产品生成完整营销内容，只输出JSON，不要其他内容：

产品名称：{product_name}
目标客户：{target_customer}
推广平台：{platform}
产品特点：{features}

JSON结构：
{{
  "商品标题": ["标题1（含关键词，25字内）", "标题2", "标题3"],
  "封面标题": ["封面文案1（强冲击力，10字内）", "封面文案2", "封面文案3", "封面文案4", "封面文案5"],
  "核心卖点": ["卖点1", "卖点2", "卖点3", "卖点4", "卖点5"],
  "小红书文案": "完整小红书图文文案，含钩子开场+产品介绍+使用场景+行动号召，800字，加emoji",
  "短视频口播稿": "完整短视频口播稿，适合30-60秒，含开场+卖点+结尾引导",
  "亚马逊卖点": ["• Feature 1 with details (英文)", "• Feature 2", "• Feature 3", "• Feature 4", "• Feature 5"]
}}"""

    response = client.chat.completions.create(
        model=model_name,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.8
    )

    raw = response.choices[0].message.content.strip()

    if "```" in raw:
        for part in raw.split("```"):
            part = part.strip().lstrip("json").strip()
            if part.startswith("{"):
                raw = part
                break

    try:
        data = json.loads(raw)
    except json.JSONDecodeError:
        data = {
            "商品标题": [], "封面标题": [], "核心卖点": [],
            "小红书文案": raw, "短视频口播稿": "", "亚马逊卖点": []
        }

    return data, response.usage


# =========================
# 批量Excel生成
# =========================

def generate_excel(results_list, filename="marketing_output.xlsx"):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "营销内容"

    columns = ["产品名称", "目标客户", "平台", "商品标题", "封面标题",
               "核心卖点", "小红书文案", "短视频口播稿", "亚马逊卖点", "生成时间"]

    header_fill = PatternFill(start_color="FF6B6B", end_color="FF6B6B", fill_type="solid")

    for col_idx, col_name in enumerate(columns, 1):
        cell = ws.cell(row=1, column=col_idx, value=col_name)
        cell.fill = header_fill
        cell.font = Font(bold=True, color="FFFFFF", size=11)
        cell.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 24

    def to_str(v):
        if isinstance(v, list):
            return "\n".join(str(i) for i in v)
        return str(v) if v else ""

    for row_idx, item in enumerate(results_list, 2):
        c = item.get("content", {})
        vals = [
            item.get("product_name", ""), item.get("target_customer", ""),
            item.get("platform", ""),
            to_str(c.get("商品标题", "")), to_str(c.get("封面标题", "")),
            to_str(c.get("核心卖点", "")), to_str(c.get("小红书文案", "")),
            to_str(c.get("短视频口播稿", "")), to_str(c.get("亚马逊卖点", "")),
            item.get("time", "")
        ]
        for col_idx, val in enumerate(vals, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=val)
            cell.alignment = Alignment(wrap_text=True, vertical="top")

    for col_idx, width in enumerate([18, 12, 10, 30, 25, 25, 50, 40, 35, 18], 1):
        ws.column_dimensions[ws.cell(row=1, column=col_idx).column_letter].width = width

    wb.save(filename)
    return filename


# =========================
# Sidebar
# =========================

with st.sidebar:
    st.title("⚙️ 系统设置")

    mode = st.radio("运行模式", ["智能模式", "批量模式"])
    model_name = st.selectbox("生成模型", ["deepseek-chat", "deepseek-reasoner"])

    if mode == "智能模式":
        st.markdown("---")
        auto_feishu = st.toggle("☁️ 自动保存到飞书", value=True)

    st.markdown("---")
    st.markdown("### 📋 历史记录")
    if not st.session_state.history:
        st.caption("暂无记录")
    for item in reversed(st.session_state.history[-8:]):
        st.caption(item)


# =========================
# 主界面
# =========================

st.title("🚀 AI Marketing Copilot V3")
st.caption("分字段结构化输出 · 小红书封面生成 · 自动保存飞书")

# ===== 批量模式 =====
if mode == "批量模式":
    st.subheader("📦 批量生成")
    st.info("CSV表头：`产品名称,目标客户,推广平台,产品特点`")

    batch_file = st.file_uploader("上传产品CSV", type=["csv"])

    if batch_file:
        import csv, io
        rows = list(csv.DictReader(io.StringIO(batch_file.read().decode("utf-8-sig"))))
        st.success(f"已读取 {len(rows)} 条产品数据")

        if st.button("🚀 开始批量生成", use_container_width=True, type="primary"):
            client_ai = OpenAI(api_key=st.secrets["DEEPSEEK_API_KEY"], base_url="https://api.deepseek.com")
            results = []
            progress = st.progress(0)
            status = st.empty()

            for i, row in enumerate(rows):
                pname = row.get("产品名称", "").strip()
                status.text(f"生成中 {i+1}/{len(rows)}：{pname}")
                try:
                    content, _ = generate_structured_content(
                        pname, row.get("目标客户", "").strip(),
                        row.get("推广平台", "小红书").strip(),
                        row.get("产品特点", "").strip(),
                        client_ai, model_name
                    )
                except Exception as e:
                    content = {"小红书文案": f"生成失败：{e}"}

                results.append({
                    "product_name": pname,
                    "target_customer": row.get("目标客户", ""),
                    "platform": row.get("推广平台", "小红书"),
                    "content": content,
                    "time": datetime.now().strftime("%Y-%m-%d %H:%M")
                })
                progress.progress((i + 1) / len(rows))

            status.text("✅ 批量生成完成！")
            excel_path = generate_excel(results)
            with open(excel_path, "rb") as f:
                st.download_button(
                    "📊 下载批量Excel", f,
                    file_name="batch_marketing.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True
                )

# ===== 智能模式 =====
else:
    col1, col2 = st.columns([1, 1])

    with col1:
        uploaded_file = st.file_uploader("📷 上传产品图片（可选）", type=["jpg", "jpeg", "png"])
        image_path = None
        if uploaded_file:
            st.image(uploaded_file, caption="产品图片预览", width=280)
            suffix = uploaded_file.name.split(".")[-1]
            with tempfile.NamedTemporaryFile(delete=False, suffix=f".{suffix}") as tmp:
                tmp.write(uploaded_file.getvalue())
                image_path = tmp.name

    with col2:
        product_name = st.text_input("产品名称 *", placeholder="如：316不锈钢真空保温杯500ml")
        target_customer = st.text_input("目标客户", placeholder="如：25-35岁上班族女性")
        platform = st.selectbox("推广平台", ["小红书", "抖音", "Temu", "亚马逊"])

    features = st.text_area(
        "产品特点",
        height=120,
        placeholder="316不锈钢\n保温24小时\n防漏设计\n500ml大容量"
    )

    # 封面图区域
    st.markdown("---")
    st.markdown("#### 🖼️ 小红书封面图生成")

    cover_col1, cover_col2 = st.columns([1, 2])

    with cover_col1:
        template_name = st.selectbox("选择场景模板", ["自定义"] + list(COVER_TEMPLATES.keys()))

    with cover_col2:
        if template_name == "自定义":
            default_prompt = ""
        else:
            default_prompt = COVER_TEMPLATES[template_name].format(
                product=product_name or "产品"
            )

        cover_prompt = st.text_area(
            "封面提示词（可修改）",
            value=default_prompt,
            height=90,
            placeholder="描述你想要的封面画面，如：白色大理石桌面上摆放着一个高颜值保温杯，旁边有几朵干花，小红书风格，清新自然"
        )

    gen_cover_btn = st.button("🎨 生成封面图", use_container_width=False)

    if gen_cover_btn:
        if not cover_prompt.strip():
            st.warning("请输入封面提示词")
        else:
            with st.spinner("🎨 生成封面图中（约20-30秒）..."):
                cover_path_preview, err = generate_cover_image(cover_prompt)
                if cover_path_preview:
                    st.session_state["cover_path"] = cover_path_preview
                    st.image(cover_path_preview, caption="AI生成封面图（768×1024）", width=280)
                    with open(cover_path_preview, "rb") as f:
                        st.download_button("⬇️ 下载封面图", f,
                                           file_name="cover.jpg", mime="image/jpeg")
                else:
                    st.error(err)

    st.markdown("---")

    if st.button("🚀 AI生成营销内容", use_container_width=True, type="primary"):
        if not product_name:
            st.warning("请输入产品名称")
            st.stop()

        client_ai = OpenAI(api_key=st.secrets["DEEPSEEK_API_KEY"], base_url="https://api.deepseek.com")

        try:
            with st.spinner("✍️ 生成营销内容中..."):
                parsed, usage = generate_structured_content(
                    product_name, target_customer, platform, features, client_ai, model_name
                )

            st.success("✅ 内容生成成功")

            tab1, tab2, tab3, tab4, tab5 = st.tabs([
                "🏷️ 标题", "📌 核心卖点", "📱 小红书文案", "🎬 短视频脚本", "🌍 亚马逊"
            ])

            with tab1:
                st.markdown("**商品标题**")
                for i, t in enumerate(parsed.get("商品标题", []), 1):
                    st.write(f"{i}. {t}")
                st.markdown("**封面标题**")
                for i, t in enumerate(parsed.get("封面标题", []), 1):
                    st.write(f"{i}. {t}")

            with tab2:
                for kp in parsed.get("核心卖点", []):
                    st.write(f"• {kp}")

            with tab3:
                st.markdown(parsed.get("小红书文案", ""))

            with tab4:
                st.markdown(parsed.get("短视频口播稿", ""))

            with tab5:
                for kp in parsed.get("亚马逊卖点", []):
                    st.write(kp)

            # Token统计
            try:
                c1, c2, c3 = st.columns(3)
                c1.metric("输入Token", usage.prompt_tokens)
                c2.metric("输出Token", usage.completion_tokens)
                c3.metric("总Token", usage.total_tokens)
            except Exception:
                pass

            # 飞书保存
            if auto_feishu:
                cover_path = st.session_state.get("cover_path")
                with st.spinner("☁️ 保存到飞书..."):
                    try:
                        feishu_resp = save_to_feishu(parsed, image_path, cover_path)
                        if feishu_resp.get("code") == 0:
                            st.success("✅ 已保存到飞书")
                        else:
                            st.warning(f"飞书保存异常 code={feishu_resp.get('code')}：{feishu_resp.get('msg')}")
                    except Exception as e:
                        st.warning(f"飞书保存失败：{e}")

            st.session_state.history.append(
                f"{datetime.now().strftime('%H:%M')} | {product_name} | {platform}"
            )

            # 导出
            st.markdown("#### 导出文件")
            col_w, col_e = st.columns(2)

            with col_w:
                doc = Document()
                doc.add_heading("AI营销方案", level=1)
                for k, v in parsed.items():
                    doc.add_heading(k, level=2)
                    doc.add_paragraph(v if isinstance(v, str) else "\n".join(str(i) for i in v))
                doc.save("marketing_plan.docx")
                with open("marketing_plan.docx", "rb") as f:
                    st.download_button("📄 下载Word", f, file_name="marketing_plan.docx")

            with col_e:
                excel_path = generate_excel([{
                    "product_name": product_name,
                    "target_customer": target_customer,
                    "platform": platform,
                    "content": parsed,
                    "time": datetime.now().strftime("%Y-%m-%d %H:%M")
                }])
                with open(excel_path, "rb") as f:
                    st.download_button(
                        "📊 下载Excel", f,
                        file_name="marketing_plan.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    )

        except Exception as e:
            st.error(f"发生错误：{e}")
            st.exception(e)